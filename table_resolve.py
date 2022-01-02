import mammoth
from docx import Document
from docx.oxml.text.paragraph import CT_P
from docx.oxml.table import CT_Tbl
from docx.table import Table
from docx.text.paragraph import Paragraph
from copy import deepcopy
import re


def write_md():
    with open("document.docx", "rb") as docx_file:
        result = mammoth.convert_to_html(docx_file)
        html = result.value  # The generated HTML
        messages = result.messages  # Any messages, such as warnings during conversion

    with open("docx_table.md", "w") as f:
        f.write(html)


def BNF_parse(paragraphs: Paragraph):
    print(paragraphs.text)
    # 正文 大写字母开头，包括- 数字 三种，如果有For example，或者包含example的单词，flag_code记下,如果有Syntax 开头，flag_syntax记下
    # BNF分析，小写字母开头，有flag标记
    # code分析，小写字母开头，有flag标记,
    # code每行分析, 缩进标记, flag_table
    text_patten = re.compile(r'^[A-Z-0-9][\s\w].*?')
    code_patten = re.compile(r'^[a-z][\s\w].*?')
    flag_code = False
    flag_syntax = False
    flag_code_table = [0, '  ']
    text = ''
    if re.match(text_patten, paragraphs.text):
        if re.match('Syntax', paragraphs.text):
            flag_syntax = True
        elif re.match(r'^For example|example', paragraphs.text):
            flag_code = True
        elif re.match(r'Syntax [0-9-].*?—Syntax', paragraphs.text):
            flag_syntax = False
        else:
            flag_code = False
            flag_syntax = False
        pass
    if flag_syntax:
        # text =  code_parse(paragraphs.text)
        pass
    if flag_code:
        # flag_code_table[0], text =  code_parse(flag_code, paragraphs.text)
        pass
    return text


def table_parse_write(table: Table):
    document = Document()
    new_table = deepcopy(table)
    paragraph = document.add_paragraph()
    paragraph._p.addnext(new_table._element)
    with open("docx_temp.docx", "wb+") as f:
        document.save(f.name)
        f.seek(0)
        result = mammoth.convert_to_html(f)
        html = result.value
        return html


def p_t():
    input_doc = Document("document.docx")
    for child in input_doc.element.body.xpath('w:p | w:tbl'):
        if isinstance(child, CT_P):
            paragraph = Paragraph(child, input_doc)
            BNF_parse(paragraph)
        elif isinstance(child, CT_Tbl):
            table = Table(child, input_doc)
            table_parse_write(table)


if __name__ == '__main__':
    # table_parse()
    p_t()
